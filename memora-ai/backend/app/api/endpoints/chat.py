from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import Any
import openai
import os
import pymupdf as fitz # pymupdf
import docx
from app.models.document import Document, DocumentChunk
from app.services.rag import get_embeddings, process_text_into_chunks

from app.api import deps
from app.models.user import User
from app.schemas.document import ChatRequest
from app.services.rag import search_similar_chunks, ask_memora
from pydantic import BaseModel

class MemoraChatRequest(BaseModel):
    message: str

from app.core.config import settings

router = APIRouter()

@router.post("/")
def chat_with_knowledge(
    *,
    db: Session = Depends(deps.get_db),
    request: ChatRequest,
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    # Get the latest user message
    user_message = request.messages[-1].content
    
    # 1. Search Vector Database for context
    try:
        similar_chunks = search_similar_chunks(db, user_message, limit=5)
    except Exception as e:
        # Fallback if pgvector fails or is not setup yet
        similar_chunks = []
        
    context_text = "\n\n".join([chunk['content'] for chunk in similar_chunks])
    
    if not context_text:
        return {
            "answer": "I couldn't find enough information in the uploaded company knowledge to answer that.",
            "citations": []
        }

    # 2. Build Prompt for LLM
    system_prompt = f"""You are an AI-Powered Enterprise Knowledge Assistant.
Your primary objective is to answer the user's question ONLY using the provided retrieved context.
If the context does not contain the answer, you must respond with exactly: "I couldn't find enough information in the uploaded company knowledge to answer that."
Do not rely on your general knowledge. Never hallucinate or invent information.

FORMAT YOUR ANSWER AS FOLLOWS:
- Summary
- Detailed Explanation
- Relevant Code Snippets (if any)
- Best Practices (if any)

RETRIEVED CONTEXT:
{context_text}
"""
    
    messages = [{"role": "system", "content": system_prompt}]
    for msg in request.messages:
        messages.append({"role": msg.role, "content": msg.content})

    # 3. Call OpenAI LLM
    if not settings.OPENAI_API_KEY:
        return {
            "answer": "Warning: OpenAI API Key is missing. Here is the context found:\n\n" + context_text,
            "citations": similar_chunks
        }
        
    client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4o-mini", # Using modern default
        messages=messages,
        temperature=0.0
    )
    
    return {
        "answer": response.choices[0].message.content,
        "citations": similar_chunks
    }

@router.post("/upload")
async def upload_document(
    *,
    db: Session = Depends(deps.get_db),
    file: UploadFile = File(...),
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    # Read file content
    content = await file.read()
    text_content = ""
    file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ""
    
    try:
        if file_extension == 'pdf':
            pdf_document = fitz.open(stream=content, filetype="pdf")
            for page in pdf_document:
                text_content += page.get_text()
        elif file_extension in ['docx', 'doc']:
            import io
            doc = docx.Document(io.BytesIO(content))
            text_content = "\n".join([para.text for para in doc.paragraphs])
        else:
            # Fallback to plain text decode
            text_content = content.decode('utf-8', errors='ignore')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")
    
    # Process text into chunks
    chunks = process_text_into_chunks(text_content)
    if not chunks:
        raise HTTPException(status_code=400, detail="Could not extract text from file")
        
    embedder = get_embeddings()
    embeddings = embedder.embed_documents(chunks)
    
    # Save document
    doc = Document(filename=file.filename, file_type=file.filename.split('.')[-1], uploader_id=current_user.id)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    # Save chunks
    for chunk_text, embedding in zip(chunks, embeddings):
        db_chunk = DocumentChunk(
            document_id=doc.id,
            content=chunk_text,
            embedding=embedding
        )
        db.add(db_chunk)
    db.commit()
    
    return {"status": "success", "filename": file.filename, "chunks_ingested": len(chunks)}

@router.post("/memora")
async def chat_with_memora(
    request: MemoraChatRequest,
    db: Session = Depends(deps.get_db)
) -> Any:
    """
    Dedicated endpoint for the VS Code extension to interact with the Two-Tier RAG system.
    """
    reply = ask_memora(db, request.message)
    return {"reply": reply}
